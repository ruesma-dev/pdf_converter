# pdf_to_word.py

import os
import io
import sys
from PIL import Image
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
import pytesseract
from tkinter import messagebox

def resource_path(relative_path):
    """Obtiene la ruta absoluta del recurso, funciona para scripts y para PyInstaller."""
    try:
        # PyInstaller crea un atributo _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Configurar ruta a Tesseract si es necesario
# Asegúrate de incluir 'tesseract.exe' en la carpeta 'assets/' y de agregarla al ejecutable
pytesseract.pytesseract.tesseract_cmd = resource_path(os.path.join('assets', 'tesseract.exe'))

def pdf_to_word(pdf_path: str, output_folder: str, use_ocr: bool = False) -> None:
    """
    Convierte un PDF a un documento Word (.docx).
    - Si use_ocr = True, aplica OCR en páginas sin texto legible.
    """
    try:
        pdf_document = fitz.open(pdf_path)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo abrir el PDF.\n{e}")
        return

    doc = Document()
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        text = page.get_text()
        images = page.get_images(full=True)

        # Si no hay texto y está activado el OCR, procesar la página como imagen
        if use_ocr and not text.strip():
            try:
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img, lang='spa')
            except Exception as e:
                messagebox.showerror("Error", f"OCR falló en la página {page_number + 1}.\n{e}")

        # Añadir texto al documento (si lo hay)
        if text.strip():
            doc.add_paragraph(text)

        # Añadir imágenes al documento
        for img_info in images:
            xref = img_info[0]
            try:
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_stream = io.BytesIO(image_bytes)
                doc.add_picture(image_stream, width=Inches(5))
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo extraer una imagen en la página {page_number + 1}.\n{e}")

        # Salto de página en Word si no es la última página
        if page_number < len(pdf_document) - 1:
            doc.add_page_break()

    output_path = os.path.join(output_folder, f"{pdf_name}.docx")
    try:
        doc.save(output_path)
        messagebox.showinfo("Proceso completado", f"Documento Word guardado en:\n{output_path}")
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo guardar el documento Word.\n{e}")
    finally:
        pdf_document.close()
