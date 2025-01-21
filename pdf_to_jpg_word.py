# pdf_to_jpg_word.py

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import os
from tkinter import Tk, filedialog, ttk, Button, StringVar, DISABLED, NORMAL, BooleanVar, Checkbutton
from PIL import Image
import io
import pytesseract

# Configurar ruta a Tesseract si es necesario
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def pdf_to_word(pdf_path, output_folder, use_ocr=False):
    """
    Convierte un archivo PDF a un documento Word manteniendo el texto y las imágenes.
    Usa OCR si el texto no es legible y la opción está activada.

    Args:
        pdf_path (str): Ruta al archivo PDF.
        output_folder (str): Carpeta para guardar el archivo Word.
        use_ocr (bool): Indica si debe usarse OCR en imágenes.

    Returns:
        None
    """
    pdf_document = fitz.open(pdf_path)
    doc = Document()

    print(f"Convirtiendo el archivo PDF {pdf_path} a Word...")

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        text = page.get_text()
        images = page.get_images(full=True)

        # Si no hay texto y está activado el OCR, procesar la página como imagen
        if use_ocr and not text.strip():
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, lang='spa')

        # Añadir texto al documento
        if text.strip():
            doc.add_paragraph(text)

        # Añadir imágenes al documento
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_stream = io.BytesIO(image_bytes)
            image = Image.open(image_stream)

            image_stream.seek(0)
            doc.add_picture(image_stream, width=Inches(5))

        if page_number < len(pdf_document) - 1:
            doc.add_page_break()

    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_path = os.path.join(output_folder, f"{pdf_name}.docx")
    doc.save(output_path)
    pdf_document.close()
    print(f"Conversión completada. Documento guardado en {output_path}")


def pdf_to_jpg(pdf_path, output_folder, dpi=300):
    """
    Convierte cada página de un archivo PDF en una imagen JPG de alta calidad.

    Args:
        pdf_path (str): Ruta al archivo PDF.
        output_folder (str): Carpeta donde se guardarán las imágenes.
        dpi (int): Resolución en DPI para las imágenes generadas. (Por defecto 300)

    Returns:
        None
    """
    pdf_document = fitz.open(pdf_path)
    total_pages = len(pdf_document)

    print(f"Convirtiendo {total_pages} páginas del PDF a imágenes...")

    for page_number in range(total_pages):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        output_path = os.path.join(output_folder, f"page_{page_number + 1}.jpg")
        img.save(output_path, "JPEG", quality=95)
        print(f"Página {page_number + 1} guardada en {output_path}")

    pdf_document.close()
    print("Conversión completada.")


if __name__ == "__main__":
    root = Tk()
    root.title("Conversor de PDF")
    root.geometry("400x250")

    file_path = StringVar()
    format_choice = StringVar(value="word")
    use_ocr = BooleanVar(value=False)

    def select_file():
        file = filedialog.askopenfilename(title="Selecciona un archivo PDF", filetypes=[("Archivos PDF", "*.pdf")])
        file_path.set(file)
        execute_button.config(state=NORMAL if file else DISABLED)

    def execute_conversion():
        pdf_path = file_path.get()
        if pdf_path:
            base_dir = os.path.dirname(pdf_path)
            pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_folder = os.path.join(base_dir, f"{pdf_name}_{format_choice.get()}")
            os.makedirs(output_folder, exist_ok=True)

            if format_choice.get() == 'word':
                pdf_to_word(pdf_path, output_folder, use_ocr.get())
            elif format_choice.get() == 'jpg':
                pdf_to_jpg(pdf_path, output_folder)

            root.destroy()

    Button(root, text="Seleccionar Archivo PDF", command=select_file).pack(pady=10)
    ttk.Radiobutton(root, text="Word", variable=format_choice, value="word").pack()
    ttk.Radiobutton(root, text="JPG", variable=format_choice, value="jpg").pack()
    Checkbutton(root, text="Usar OCR", variable=use_ocr).pack(pady=5)
    execute_button = Button(root, text="Ejecutar", command=execute_conversion, state=DISABLED)
    execute_button.pack(pady=20)

    root.mainloop()
