"""
Script unificado para gestionar:
1) Convertir imágenes a PDF
2) Convertir PDF a Word
3) Convertir PDF a JPG
4) Extraer páginas de un PDF

Requerimientos:
- Python 3.12
- Librerías:
    pip install Pillow PyMuPDF python-docx pytesseract (opcional)
"""

import os
import io
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import ttk, filedialog, simpledialog, messagebox

# -------------------------------------------------------------------
# Funciones de cada funcionalidad
# -------------------------------------------------------------------

def images_to_pdf(input_path: str, output_path: str) -> None:
    """
    Convierte una carpeta con imágenes (o una sola imagen)
    en un único archivo PDF.
    """
    valid_extensions = ['.jpg', '.jpeg', '.png', '.bmp', '.gif']
    images = []

    if os.path.isdir(input_path):
        # Obtener todas las imágenes de la carpeta
        for file in sorted(os.listdir(input_path)):
            if os.path.splitext(file)[1].lower() in valid_extensions:
                img_path = os.path.join(input_path, file)
                img = Image.open(img_path).convert('RGB')
                images.append(img)
    else:
        # Procesar una sola imagen
        if os.path.splitext(input_path)[1].lower() in valid_extensions:
            img = Image.open(input_path).convert('RGB')
            images.append(img)

    if not images:
        messagebox.showwarning("Aviso", "No se encontraron imágenes válidas para convertir.")
        return

    images[0].save(output_path, save_all=True, append_images=images[1:])
    messagebox.showinfo("Proceso completado", f"PDF guardado en:\n{output_path}")


def pdf_to_word(pdf_path: str, output_folder: str, use_ocr: bool = False) -> None:
    """
    Convierte un PDF a un documento Word (.docx).
    - Si use_ocr = True, aplica OCR en páginas sin texto legible.
    """
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        text = page.get_text()
        images = page.get_images(full=True)

        # Si no hay texto y está activado el OCR, procesar la página como imagen
        if use_ocr and not text.strip():
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img)

        # Añadir texto al documento (si lo hay)
        if text.strip():
            doc.add_paragraph(text)

        # Añadir imágenes al documento
        for _, img_info in enumerate(images):
            xref = img_info[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_stream = io.BytesIO(image_bytes)
            image = Image.open(image_stream)

            image_stream.seek(0)
            doc.add_picture(image_stream, width=Inches(5))

        # Salto de página en Word si no es la última página
        if page_number < len(pdf_document) - 1:
            doc.add_page_break()

    output_path = os.path.join(output_folder, f"{pdf_name}.docx")
    doc.save(output_path)
    pdf_document.close()
    messagebox.showinfo("Proceso completado", f"Documento Word guardado en:\n{output_path}")


def pdf_to_jpg(pdf_path: str, output_folder: str, dpi: int = 300) -> None:
    """
    Convierte cada página de un PDF en una imagen JPG.
    """
    pdf_document = fitz.open(pdf_path)
    total_pages = len(pdf_document)
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]

    for page_number in range(total_pages):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        # Nombre de salida con número de página
        output_path = os.path.join(output_folder, f"{pdf_name}_page_{page_number + 1}.jpg")
        img.save(output_path, "JPEG", quality=95)

    pdf_document.close()
    messagebox.showinfo(
        "Proceso completado",
        f"Imágenes JPG generadas en la carpeta:\n{output_folder}"
    )


def extract_pages_from_pdf(input_pdf: str, output_pdf: str, start_page: int, end_page: int) -> None:
    """
    Extrae páginas específicas de un archivo PDF y las guarda en un nuevo PDF.
    """
    pdf_document = fitz.open(input_pdf)

    if (
        start_page < 1 or
        end_page > len(pdf_document) or
        start_page > end_page
    ):
        messagebox.showwarning("Aviso", "Rango de páginas inválido.")
        pdf_document.close()
        return

    output_document = fitz.open()

    for page_num in range(start_page - 1, end_page):
        output_document.insert_pdf(pdf_document, from_page=page_num, to_page=page_num)

    output_document.save(output_pdf)
    output_document.close()
    pdf_document.close()

    messagebox.showinfo(
        "Proceso completado",
        f"Páginas {start_page}-{end_page} extraídas y guardadas en:\n{output_pdf}"
    )


# -------------------------------------------------------------------
# Interfaz gráfica unificada
# -------------------------------------------------------------------

class AppUnificada(tk.Tk):
    """
    Ventana principal que unifica las funcionalidades:
    1) Convertir imágenes a PDF
    2) Convertir PDF a Word
    3) Convertir PDF a JPG
    4) Extraer páginas de PDF
    """

    def __init__(self):
        super().__init__()
        self.title("App Unificada PDF / Imágenes")
        self.geometry("500x400")

        # Variables de control
        self.option_var = tk.StringVar(value="img_to_pdf")  # Valor por defecto
        self.file_path_var = tk.StringVar()
        self.output_path_var = tk.StringVar()
        self.ocr_var = tk.BooleanVar(value=False)
        self.dpi_var = tk.StringVar(value="300")  # Solo para PDF->JPG
        self.start_page_var = tk.StringVar()
        self.end_page_var = tk.StringVar()

        # Crear interfaz
        self.create_widgets()

    def create_widgets(self):
        """Crea los elementos de la interfaz gráfica."""
        # Frame para opciones (radio buttons)
        options_frame = tk.LabelFrame(self, text="Selecciona la funcionalidad")
        options_frame.pack(fill="x", padx=10, pady=10)

        tk.Radiobutton(
            options_frame, text="Imágenes a PDF",
            variable=self.option_var, value="img_to_pdf",
            command=self.update_ui
        ).pack(anchor="w", padx=5, pady=2)

        tk.Radiobutton(
            options_frame, text="PDF a Word",
            variable=self.option_var, value="pdf_to_word",
            command=self.update_ui
        ).pack(anchor="w", padx=5, pady=2)

        tk.Radiobutton(
            options_frame, text="PDF a JPG",
            variable=self.option_var, value="pdf_to_jpg",
            command=self.update_ui
        ).pack(anchor="w", padx=5, pady=2)

        tk.Radiobutton(
            options_frame, text="Extraer páginas de PDF",
            variable=self.option_var, value="extract_pages",
            command=self.update_ui
        ).pack(anchor="w", padx=5, pady=2)

        # Frame para seleccionar archivo(s) de entrada
        input_frame = tk.LabelFrame(self, text="Entrada")
        input_frame.pack(fill="x", padx=10, pady=5)

        tk.Label(input_frame, text="Seleccionar archivo o carpeta").grid(row=0, column=0, padx=5, pady=5)
        tk.Entry(input_frame, textvariable=self.file_path_var, width=40).grid(row=0, column=1, padx=5, pady=5)
        tk.Button(input_frame, text="Examinar", command=self.browse_input).grid(row=0, column=2, padx=5, pady=5)

        # Frame para opciones específicas
        self.options_frame = tk.LabelFrame(self, text="Opciones adicionales")
        self.options_frame.pack(fill="x", padx=10, pady=5)

        # 1) OCR
        self.ocr_check = tk.Checkbutton(
            self.options_frame, text="Usar OCR (para PDF->Word)",
            variable=self.ocr_var
        )
        self.ocr_check.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        # 2) DPI
        tk.Label(self.options_frame, text="DPI (para PDF->JPG):").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        tk.Entry(self.options_frame, textvariable=self.dpi_var, width=6).grid(row=1, column=1, padx=5, pady=5, sticky="w")

        # 3) Rango de páginas (para extraer páginas)
        tk.Label(self.options_frame, text="Página inicial:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        tk.Entry(self.options_frame, textvariable=self.start_page_var, width=6).grid(row=2, column=1, padx=5, pady=5, sticky="w")

        tk.Label(self.options_frame, text="Página final:").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        tk.Entry(self.options_frame, textvariable=self.end_page_var, width=6).grid(row=3, column=1, padx=5, pady=5, sticky="w")

        # Frame para la salida (PDF, carpeta, etc.)
        output_frame = tk.LabelFrame(self, text="Salida")
        output_frame.pack(fill="x", padx=10, pady=5)

        tk.Label(output_frame, text="Archivo/Carpeta destino").grid(row=0, column=0, padx=5, pady=5)
        tk.Entry(output_frame, textvariable=self.output_path_var, width=40).grid(row=0, column=1, padx=5, pady=5)
        tk.Button(output_frame, text="Guardar en...", command=self.browse_output).grid(row=0, column=2, padx=5, pady=5)

        # Botón ejecutar
        tk.Button(self, text="Ejecutar", command=self.execute_action, bg="green", fg="white").pack(pady=10)

        self.update_ui()

    def browse_input(self):
        """Selecciona archivo(s) o carpeta según la opción activa."""
        option = self.option_var.get()

        if option == "img_to_pdf":
            # Primero intentamos seleccionar una carpeta; si no, seleccionamos una imagen
            folder = filedialog.askdirectory(title="Selecciona una carpeta de imágenes")
            if folder:
                self.file_path_var.set(folder)
            else:
                # Seleccionar archivo de imagen
                filetypes = [("Imágenes", "*.jpg;*.jpeg;*.png;*.bmp;*.gif")]
                file = filedialog.askopenfilename(title="Selecciona una imagen", filetypes=filetypes)
                if file:
                    self.file_path_var.set(file)
        else:
            # En las otras opciones, seleccionamos un PDF
            filetypes = [("Archivos PDF", "*.pdf")]
            file = filedialog.askopenfilename(title="Selecciona un archivo PDF", filetypes=filetypes)
            if file:
                self.file_path_var.set(file)

    def browse_output(self):
        """Selecciona la ruta de salida, que puede ser archivo o carpeta."""
        option = self.option_var.get()

        if option == "img_to_pdf":
            # Guardar como PDF
            file = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("Archivos PDF", "*.pdf")],
                title="Guardar PDF"
            )
            if file:
                self.output_path_var.set(file)
        elif option == "extract_pages":
            # También guardar como PDF
            file = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("Archivos PDF", "*.pdf")],
                title="Guardar PDF"
            )
            if file:
                self.output_path_var.set(file)
        else:
            # PDF -> Word o PDF -> JPG => Seleccionar carpeta de salida
            folder = filedialog.askdirectory(title="Selecciona carpeta de salida")
            if folder:
                self.output_path_var.set(folder)

    def update_ui(self):
        """Muestra u oculta las opciones relevantes según la funcionalidad."""
        option = self.option_var.get()

        # Resetear visibilidad
        self.ocr_check.grid_remove()
        for idx in range(1, 4):
            # Ocultar filas 1 a 3 (DPI y rango de páginas)
            self.options_frame.grid_rowconfigure(idx, minsize=0)
            for widget in self.options_frame.grid_slaves(row=idx):
                widget.grid_remove()

        # Imágenes a PDF => no necesitamos OCR, DPI, ni rango de páginas
        if option == "img_to_pdf":
            pass  # Sin campos extra

        # PDF a Word => activar opción de OCR
        elif option == "pdf_to_word":
            self.ocr_check.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        # PDF a JPG => mostrar DPI
        elif option == "pdf_to_jpg":
            # OCR no aplica, pero DPI sí
            # Reubicar DPI en la fila 1
            self.options_frame.grid_rowconfigure(1, minsize=30)
            for widget in self.options_frame.grid_slaves(row=1):
                widget.grid()

        # Extraer páginas => mostrar inputs para start_page y end_page
        elif option == "extract_pages":
            # Filas 2 y 3 visibles
            self.options_frame.grid_rowconfigure(2, minsize=30)
            self.options_frame.grid_rowconfigure(3, minsize=30)
            for widget in self.options_frame.grid_slaves(row=2):
                widget.grid()
            for widget in self.options_frame.grid_slaves(row=3):
                widget.grid()

    def execute_action(self):
        """Ejecuta la funcionalidad seleccionada."""
        option = self.option_var.get()
        input_path = self.file_path_var.get().strip()
        output_path = self.output_path_var.get().strip()

        if not input_path:
            messagebox.showwarning("Aviso", "Por favor, selecciona la ruta de entrada.")
            return
        if option == "img_to_pdf" or option == "extract_pages":
            if not output_path:
                messagebox.showwarning("Aviso", "Por favor, selecciona la ruta de salida.")
                return
        else:
            # PDF->Word o PDF->JPG => se usa una carpeta de salida
            if not output_path or not os.path.isdir(output_path):
                messagebox.showwarning("Aviso", "Por favor, selecciona una carpeta de salida válida.")
                return

        try:
            if option == "img_to_pdf":
                images_to_pdf(input_path, output_path)

            elif option == "pdf_to_word":
                use_ocr = self.ocr_var.get()
                pdf_to_word(input_path, output_path, use_ocr)

            elif option == "pdf_to_jpg":
                dpi_value = int(self.dpi_var.get())
                pdf_to_jpg(input_path, output_path, dpi=dpi_value)

            elif option == "extract_pages":
                start_page_str = self.start_page_var.get().strip()
                end_page_str = self.end_page_var.get().strip()

                # Verificar que ambos campos tengan valores y que sean dígitos
                if not (start_page_str.isdigit() and end_page_str.isdigit()):
                    messagebox.showwarning(
                        "Aviso",
                        "Por favor, introduce números válidos en 'Página inicial' y 'Página final'."
                    )
                    return

                start_page = int(start_page_str)
                end_page = int(end_page_str)

                if start_page < 1 or end_page < 1:
                    messagebox.showwarning(
                        "Aviso",
                        "Las páginas deben ser números enteros positivos."
                    )
                    return

                if start_page > end_page:
                    messagebox.showwarning(
                        "Aviso",
                        "La página inicial no puede ser mayor que la página final."
                    )
                    return

                extract_pages_from_pdf(input_path, output_path, start_page, end_page)

        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error:\n{e}")


if __name__ == "__main__":
    app = AppUnificada()
    app.mainloop()
